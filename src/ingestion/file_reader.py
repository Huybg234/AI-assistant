import io
import os
import re
from typing import List, Union

import PyPDF2

from src.utils import MIME_TO_EXT, SUPPORTED_EXTENSIONS


def extract_pages_from_pdf(pdf_source: Union[str, bytes]) -> List[str]:
    """
    Extracts text from each page in the PDF file.
    The source can be a file path (str) or file content (bytes).
    """
    pages = []
    try:
        if isinstance(pdf_source, str):
            file_stream = open(pdf_source, "rb")
        elif isinstance(pdf_source, bytes):
            file_stream = io.BytesIO(pdf_source)
        else:
            raise TypeError("PDF source must be a file path (str) or bytes.")

        reader = PyPDF2.PdfReader(file_stream)
        for page in reader.pages:
            page_text = page.extract_text()
            pages.append(page_text if page_text else "")
    finally:
        if 'file_stream' in locals() and hasattr(file_stream, 'close'):
            file_stream.close()

    return pages


def extract_pages_from_docx(source: Union[str, bytes]) -> List[str]:
    """Extract text from a DOCX file. Returns list of paragraph groups as 'pages'."""
    import docx as _docx

    if isinstance(source, bytes):
        doc = _docx.Document(io.BytesIO(source))
    else:
        doc = _docx.Document(source)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    page_size = 30
    pages = []
    for i in range(0, max(len(paragraphs), 1), page_size):
        page_text = "\n".join(paragraphs[i:i + page_size])
        if page_text.strip():
            pages.append(page_text)
    return pages if pages else [""]


def _clean_doc_text(text: str) -> str:
    """
    Filter out OLE metadata garbage from DOC binary text extraction.
    Removes: OLE object refs, font/style table names, binary character sequences.
    """
    _STYLE_NAMES = {
        "Normal", "Default", "Default Style", "Heading", "Title", "Subtitle",
        "Caption", "Preformatted Text", "Text Body", "Table Contents",
        "Table Heading", "Quotations", "Index", "Footer", "Header",
        "Internet Link", "Visited Internet Link", "Bullets", "List",
        "List Paragraph", "Body Text", "Block Text", "Balloon Text",
        "No Spacing", "Light Shading", "Light List", "Light Grid",
    }

    _FONT_NAMES = {
        "Symbol", "Wingdings", "Wingdings 2", "Wingdings 3", "Webdings",
        "Marlett", "Times New Roman", "Arial", "Courier New", "Tahoma",
        "Verdana", "Georgia", "Calibri", "Cambria", "Palatino",
        "DejaVu Sans", "DejaVu Serif", "DejaVu Sans Mono",
        "Liberation Serif", "Liberation Sans", "Liberation Mono",
        "Open Sans", "Droid Sans", "Droid Sans Fallback",
        "FreeSans", "FreeMono", "FreeSerif",
        "OpenSymbol", "Arial Unicode MS",
    }

    clean_lines = []
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue

        if s.startswith("EMBED ") or s.startswith("LINK ") or "LibreOffice." in s or "Microsoft.Office" in s:
            continue

        if s in _STYLE_NAMES or s in _FONT_NAMES:
            continue

        if re.match(r'^(Heading|List|Index|Table|Outline)\s+\d+$', s, re.I):
            continue

        alnum = sum(1 for c in s if c.isalnum() or c.isspace())
        if len(s) > 4 and alnum < len(s) * 0.40:
            continue

        if len(s) <= 3 and not any(c.isalpha() for c in s):
            continue

        if " " not in s and len(s) > 6:
            special_count = sum(1 for c in s if not c.isalnum())
            if special_count >= 3:
                continue

        clean_lines.append(s)

    return "\n".join(clean_lines)


def _extract_text_from_doc_ole(data: bytes) -> str:
    """
    Extract plain text from a binary .doc file using olefile.
    Only reads the WordDocument stream (actual text).
    0Table/1Table are skipped - they contain font/style metadata, not text.
    """
    import olefile

    collected: List[str] = []

    try:
        ole = olefile.OleFileIO(io.BytesIO(data))
        if ole.exists("WordDocument"):
            raw = ole.openstream("WordDocument").read()
            for match in re.finditer(b'(?:[\x20-\x7e\x09\x0a\x0d]\x00){6,}', raw):
                try:
                    seg = match.group(0).decode("utf-16-le", errors="ignore").strip()
                    if seg and len(seg) > 4:
                        collected.append(seg)
                except Exception:
                    pass
        ole.close()
    except Exception:
        pass

    if not collected:
        for match in re.finditer(rb'[ -~\t\r\n]{8,}', data):
            s = match.group(0).decode("ascii", errors="ignore").strip()
            if s and len(s) > 6:
                collected.append(s)

    raw_text = "\n".join(collected)
    return _clean_doc_text(raw_text)


def extract_pages_from_doc(source: Union[str, bytes]) -> List[str]:
    """Extract text from a legacy DOC file.
    Tries mammoth first (handles DOCX-labelled-as-.doc), then falls back to
    OLE binary extraction for true binary .doc files.
    """
    import mammoth

    if isinstance(source, bytes):
        data = source
    else:
        with open(source, "rb") as f:
            data = f.read()

    text = ""
    try:
        result = mammoth.extract_raw_text(io.BytesIO(data))
        text = (result.value or "").strip()
    except Exception:
        pass

    if not text:
        text = _extract_text_from_doc_ole(data).strip()

    if not text:
        return [""]

    lines = [line for line in text.splitlines() if line.strip()]
    page_size = 50
    pages = []
    for i in range(0, max(len(lines), 1), page_size):
        page_text = "\n".join(lines[i:i + page_size])
        if page_text.strip():
            pages.append(page_text)
    return pages if pages else [""]


def extract_pages_from_txt(source: Union[str, bytes], encoding: str = "utf-8") -> List[str]:
    """Extract text from a plain-text file."""
    if isinstance(source, bytes):
        try:
            text = source.decode(encoding)
        except UnicodeDecodeError:
            text = source.decode("latin-1", errors="replace")
    else:
        with open(source, "r", encoding=encoding, errors="replace") as f:
            text = f.read()
    lines = [line for line in text.splitlines() if line.strip()]
    page_size = 60
    pages = []
    for i in range(0, max(len(lines), 1), page_size):
        page_text = "\n".join(lines[i:i + page_size])
        if page_text.strip():
            pages.append(page_text)
    return pages if pages else [""]


def extract_pages_from_rtf(source: Union[str, bytes]) -> List[str]:
    """Extract text from an RTF file using striprtf."""
    from striprtf.striprtf import rtf_to_text

    if isinstance(source, bytes):
        try:
            rtf_str = source.decode("utf-8")
        except UnicodeDecodeError:
            rtf_str = source.decode("latin-1", errors="replace")
    else:
        with open(source, "r", encoding="latin-1", errors="replace") as f:
            rtf_str = f.read()
    text = rtf_to_text(rtf_str).strip()
    if not text:
        return [""]
    lines = [line for line in text.splitlines() if line.strip()]
    page_size = 50
    pages = []
    for i in range(0, max(len(lines), 1), page_size):
        page_text = "\n".join(lines[i:i + page_size])
        if page_text.strip():
            pages.append(page_text)
    return pages if pages else [""]


def extract_pages_from_file(
    source: Union[str, bytes],
    filename: str = "",
    content_type: str = "",
) -> List[str]:
    """
    Generic dispatcher: extract pages/sections of text from any supported file.
    Detects format from filename extension or content_type.
    """
    ext = ""
    if filename:
        ext = os.path.splitext(filename.lower())[1]
    if not ext and content_type:
        ext = MIME_TO_EXT.get(content_type, "")

    if ext == ".pdf":
        return extract_pages_from_pdf(source)
    if ext == ".docx":
        return extract_pages_from_docx(source)
    if ext == ".doc":
        return extract_pages_from_doc(source)
    if ext == ".txt":
        return extract_pages_from_txt(source)
    if ext == ".rtf":
        return extract_pages_from_rtf(source)
    raise ValueError(
        f"Unsupported file format '{ext or content_type}'. "
        f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )
